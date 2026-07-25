from pathlib import Path
from docx import Document

from app.schema.DocumentModel import (
    DocumentModel,
    DocumentElements,
    type_format,
)


def export_docx(document: DocumentModel, output_path: str) -> str:
    """
    Convert DocumentModel -> docx

    Returns file path.
    """

    doc = Document()

    # setting the metadata from the documentModel

    if document.metadata:
        props = doc.core_properties

        props.author = document.metadata.creator
        props.language = document.metadata.language

    # according to the pages no.

    total_pages = len(document.pages)

    for page_index, page in enumerate(document.pages):

        # sort by reading order
        elements = sorted(
            page.elements,
            key=lambda e: e.reading_order
        )

        for element in elements:
            write_element(doc, element)

        # don't insert page break after last page
        if page_index != total_pages - 1:
            doc.add_page_break()

    Path(output_path).parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    doc.save(output_path)

    return output_path


def write_element(doc: Document, element: DocumentElements):

    match element.type:

        case type_format.HEADING:
            write_heading(doc, element)

        case type_format.PARAGRAPH:
            write_paragraph(doc, element)

        case type_format.BULLETED_LIST:
            write_bulleted_list(doc, element)

        case type_format.NUMBERED_LIST:
            write_numbered_list(doc, element)

        case type_format.QUOTE:
            write_quote(doc, element)

        case type_format.CODE:
            write_code(doc, element)

        case type_format.TABLE:
            write_table(doc, element)

        case type_format.IMAGE:
            write_image(doc, element)

        case _:
            write_paragraph(doc, element)


def write_heading(doc: Document, element: DocumentElements):

    doc.add_heading(
        element.text,
        level=1,
    )


def write_paragraph(doc: Document, element: DocumentElements):

    doc.add_paragraph(element.text)


def write_bulleted_list(doc: Document, element: DocumentElements):

    for line in element.text.split("\n"):
        if line.strip():
            doc.add_paragraph(
                line,
                style="List Bullet",
            )


def write_numbered_list(doc: Document, element: DocumentElements):

    for line in element.text.split("\n"):
        if line.strip():
            doc.add_paragraph(
                line,
                style="List Number",
            )


def write_quote(doc: Document, element: DocumentElements):

    doc.add_paragraph(
        element.text,
        style="Quote",
    )


def write_code(doc: Document, element: DocumentElements):

    p = doc.add_paragraph()

    run = p.add_run(element.text)

    run.font.name = "Consolas"


def write_table(doc: Document, element: DocumentElements):

    """
    future if wanna include tables (rows/columns)
    """

    doc.add_paragraph("[TABLE]")
    doc.add_paragraph(element.text)


def write_image(doc: Document, element: DocumentElements):

    doc.add_paragraph("[IMAGE]")